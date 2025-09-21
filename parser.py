import fitz  # PyMuPDF
import docx
import re
from typing import Dict, List, Any
from pathlib import Path
import streamlit as st

class ResumeParser:
    """Resume parser optimized for HuggingFace LLM processing"""
    
    def __init__(self):
        # Initialize parser for HuggingFace system
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text("text")
                text += page_text + "\n"
            doc.close()
            return text.strip()
        except Exception as e:
            st.error(f"PDF extraction error: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"DOCX extraction error: {e}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read with errors ignored
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            st.error(f"Text file error: {e}")
            return ""

    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(str(file_path))
        elif extension == '.txt':
            return self.extract_text_from_txt(str(file_path))
        else:
            st.error(f"Unsupported file format: {extension}")
            return ""

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text for HuggingFace processing"""
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Max 2 consecutive line breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\r\n', '\n', text)  # Normalize line endings
        
        # Remove page numbers and common artifacts
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Skip likely artifacts
            if self._is_likely_artifact(line):
                continue
                
            # Skip very short lines (likely artifacts)
            if len(line) < 2:
                continue
                
            cleaned_lines.append(line)
        
        # Join lines back
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove extra spaces
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text.strip()

    def _is_likely_artifact(self, line: str) -> bool:
        """Check if a line is likely an artifact (headers, footers, page numbers)"""
        line = line.strip().lower()
        
        # Empty or very short
        if len(line) < 2:
            return True
        
        # Just numbers (page numbers)
        if re.match(r'^\d+$', line):
            return True
        
        # Page indicators
        if re.match(r'^page\s*\d+', line):
            return True
        
        # Common header/footer patterns
        artifacts = [
            'confidential', 'resume', 'curriculum vitae', 'cv',
            'personal and confidential', 'page', 'continued',
            'references available upon request'
        ]
        
        for artifact in artifacts:
            if line == artifact:
                return True
        
        # Lines with only special characters
        if re.match(r'^[^\w\s]*$', line):
            return True
        
        return False

    def extract_basic_info(self, text: str) -> Dict[str, Any]:
        """Extract basic information for fallback when HuggingFace fails"""
        info = {
            'email': '',
            'phone': '',
            'skills_mentioned': [],
            'experience_indicators': [],
            'education_indicators': []
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        info['email'] = emails[0] if emails else ''
        
        # Extract phone numbers
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # US format
            r'\b\(\d{3}\)\s*\d{3}[-.]?\d{4}\b',  # (123) 456-7890
            r'\b\+\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b'  # International
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                info['phone'] = phones[0]
                break
        
        # Extract skills (basic keyword matching)
        skill_keywords = [
            'python', 'java', 'javascript', 'c++', 'c#', 'sql', 'html', 'css',
            'react', 'angular', 'vue', 'node.js', 'django', 'flask',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git',
            'machine learning', 'data science', 'ai', 'tensorflow', 'pytorch'
        ]
        
        text_lower = text.lower()
        for skill in skill_keywords:
            if skill in text_lower:
                info['skills_mentioned'].append(skill.title())
        
        # Extract experience indicators
        exp_patterns = [
            r'(\d+)\s*(?:\+)?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'experience[:\s]*(\d+)\s*(?:years?|yrs?)',
            r'(\d+)\s*(?:years?|yrs?)\s*(?:in|with|of)'
        ]
        
        for pattern in exp_patterns:
            matches = re.findall(pattern, text_lower)
            info['experience_indicators'].extend(matches)
        
        # Extract education indicators
        education_keywords = [
            'bachelor', 'master', 'phd', 'mba', 'degree', 'university', 'college'
        ]
        
        for keyword in education_keywords:
            if keyword in text_lower:
                info['education_indicators'].append(keyword.title())
        
        return info

    def validate_text_quality(self, text: str) -> Dict[str, Any]:
        """Validate extracted text quality"""
        if not text:
            return {
                'valid': False,
                'issues': ['No text extracted'],
                'word_count': 0,
                'char_count': 0
            }
        
        word_count = len(text.split())
        char_count = len(text)
        issues = []
        
        # Check minimum content
        if word_count < 50:
            issues.append('Very short content (less than 50 words)')
        
        if char_count < 300:
            issues.append('Very short text (less than 300 characters)')
        
        # Check for excessive repetition
        words = text.lower().split()
        unique_words = len(set(words))
        if word_count > 0 and unique_words / word_count < 0.3:
            issues.append('High repetition detected')
        
        # Check for reasonable character distribution
        alpha_chars = sum(1 for c in text if c.isalpha())
        if char_count > 0 and alpha_chars / char_count < 0.5:
            issues.append('Low alphabetic content ratio')
        
        return {
            'valid': len(issues) == 0,
            'issues': issues,
            'word_count': word_count,
            'char_count': char_count,
            'unique_word_ratio': unique_words / word_count if word_count > 0 else 0
        }

    def parse_resume(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Main resume parsing function - optimized for HuggingFace LLM processing"""
        try:
            # Check file format
            file_path = Path(file_path)
            if file_path.suffix.lower() not in self.supported_formats:
                return {
                    "success": False, 
                    "error": f"Unsupported file format: {file_path.suffix}"
                }
            
            # Extract raw text
            raw_text = self.extract_text(str(file_path))
            
            if not raw_text.strip():
                return {
                    "success": False, 
                    "error": "Could not extract text from resume"
                }
            
            # Clean the text for HuggingFace processing
            cleaned_text = self.clean_text(raw_text)
            
            # Validate text quality
            quality_check = self.validate_text_quality(cleaned_text)
            
            # Extract basic fallback info
            basic_info = self.extract_basic_info(cleaned_text)
            
            # Prepare optimized text for HuggingFace (limit length for free tier)
            optimized_text = self._optimize_for_huggingface(cleaned_text)
            
            return {
                'success': True,
                'filename': filename,
                'cleaned_text': optimized_text,  # This is what HuggingFace LLM will process
                'raw_text': raw_text,
                'basic_info': basic_info,  # Fallback if LLM fails
                'quality_check': quality_check,
                'word_count': quality_check['word_count'],
                'char_count': len(optimized_text),
                'extraction_method': 'PyMuPDF/python-docx',
                'huggingface_ready': True,
                'processing_notes': []
            }
            
        except Exception as e:
            st.error(f"Resume parsing error for {filename}: {e}")
            return {
                "success": False, 
                "error": f"Failed to parse resume: {str(e)}",
                "filename": filename
            }

    def _optimize_for_huggingface(self, text: str) -> str:
        """Optimize text for HuggingFace free tier processing"""
        # For free tier, we need to be conservative with token usage
        # Mistral 7B can handle ~4000 tokens, but we want to be safe
        
        # Target: ~2000 characters (~500 tokens) for data extraction
        # This leaves room for the prompt and response
        
        if len(text) <= 2000:
            return text
        
        # If text is too long, intelligently truncate
        paragraphs = text.split('\n\n')
        
        # Prioritize certain sections
        priority_keywords = [
            'experience', 'work', 'employment', 'position', 'role',
            'education', 'degree', 'university', 'college',
            'skills', 'technical', 'programming', 'software',
            'project', 'achievement', 'accomplishment'
        ]
        
        # Score paragraphs by importance
        scored_paragraphs = []
        for para in paragraphs:
            if len(para.strip()) < 20:  # Skip very short paragraphs
                continue
                
            score = 0
            para_lower = para.lower()
            
            # Score based on keywords
            for keyword in priority_keywords:
                if keyword in para_lower:
                    score += 1
            
            # Boost score for paragraphs with numbers (dates, years)
            if re.search(r'\d{4}', para):  # Year patterns
                score += 1
            
            # Boost score for paragraphs with technical terms
            if re.search(r'\b(?:developed|managed|led|created|implemented|designed)\b', para_lower):
                score += 1
            
            scored_paragraphs.append((score, para))
        
        # Sort by score (descending) and select top paragraphs
        scored_paragraphs.sort(key=lambda x: x[0], reverse=True)
        
        selected_text = ""
        for score, para in scored_paragraphs:
            if len(selected_text + para) <= 1800:  # Leave room for safety
                selected_text += para + "\n\n"
            else:
                break
        
        if not selected_text.strip():
            # Fallback: take first 1800 characters
            selected_text = text[:1800]
        
        return selected_text.strip()

    def get_parser_stats(self) -> Dict[str, Any]:
        """Get parser statistics and capabilities"""
        return {
            'supported_formats': self.supported_formats,
            'optimization': 'HuggingFace Free Tier',
            'max_chars': 2000,
            'features': [
                'PDF text extraction (PyMuPDF)',
                'DOCX text extraction (python-docx)', 
                'Text cleaning and normalization',
                'Artifact removal (headers, footers)',
                'Intelligent text truncation',
                'Priority-based content selection',
                'Fallback info extraction',
                'Quality validation'
            ],
            'huggingface_optimized': True,
            'free_tier_friendly': True
        }

# Alias for compatibility with other systems
EnhancedResumeParser = ResumeParser

# Example usage for testing
if __name__ == "__main__":
    parser = ResumeParser()
    
    # Test with a sample file (replace with actual file path)
    # result = parser.parse_resume("sample_resume.pdf", "sample_resume.pdf")
    # print(json.dumps(result, indent=2))
    
    # Print parser capabilities
    import json
    print("Parser Statistics:")
    print(json.dumps(parser.get_parser_stats(), indent=2))